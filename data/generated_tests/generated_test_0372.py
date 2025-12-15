import re
import glob
from docx import Document
import os
import unittest
import tempfile

def task_func(directory_path: str) -> int:
    """
    Processes all Word (.docx) files in the provided directory, searching for double quotes in the text 
    and adding a backslash before each double quote to "protect" it.
    
    Parameters:
    - directory_path (str): Path to the directory containing .docx files to be processed.
    
    Returns:
    - int: Number of .docx files processed.

    Requirements:
    - re
    - docx
    - glob

    Example:
    >>> import tempfile
    >>> temp_dir = tempfile.mkdtemp()
    >>> doc = Document()
    >>> _ = doc.add_paragraph("This is a sample text with double quotes.")
    >>> doc.save(temp_dir + '/sample.docx')
    >>> task_func(temp_dir)
    1
    """

    docx_files = glob.glob(directory_path + '/*.docx')
    processed_files = 0

    for docx_file in docx_files:
        document = Document(docx_file)

        for paragraph in document.paragraphs:
            paragraph.text = re.sub(r'(?<!\\)"', r'\"', paragraph.text)

        document.save(docx_file)
        processed_files += 1

    return processed_files

class TestTaskFunc(unittest.TestCase):

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        for file in glob.glob(os.path.join(self.temp_dir, "*.docx")):
            os.remove(file)
        os.rmdir(self.temp_dir)

    def test_single_docx_file_with_quotes(self):
        doc = Document()
        _ = doc.add_paragraph('This is a "test" paragraph.')
        doc.save(os.path.join(self.temp_dir, 'test1.docx'))

        result = task_func(self.temp_dir)
        self.assertEqual(result, 1)

        # Verifying the content of the processed file
        document = Document(os.path.join(self.temp_dir, 'test1.docx'))
        self.assertEqual(document.paragraphs[0].text, 'This is a \"test\" paragraph.')

    def test_no_docx_files(self):
        result = task_func(self.temp_dir)
        self.assertEqual(result, 0)

    def test_multiple_docx_files_with_quotes(self):
        for i in range(3):
            doc = Document()
            _ = doc.add_paragraph(f'This is a "test" in file {i}.')
            doc.save(os.path.join(self.temp_dir, f'test{i}.docx'))

        result = task_func(self.temp_dir)
        self.assertEqual(result, 3)

        for i in range(3):
            document = Document(os.path.join(self.temp_dir, f'test{i}.docx'))
            self.assertEqual(document.paragraphs[0].text, f'This is a \"test\" in file {i}.')

    def test_docx_file_without_quotes(self):
        doc = Document()
        _ = doc.add_paragraph('This paragraph has no quotes.')
        doc.save(os.path.join(self.temp_dir, 'test2.docx'))

        result = task_func(self.temp_dir)
        self.assertEqual(result, 1)

        # Verifying the content remains unchanged
        document = Document(os.path.join(self.temp_dir, 'test2.docx'))
        self.assertEqual(document.paragraphs[0].text, 'This paragraph has no quotes.')

if __name__ == '__main__':
    unittest.main()